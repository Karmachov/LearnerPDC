# --- Standard Library Imports ---
import os
import tempfile
import platform
import time
import getpass
import traceback
import warnings
import re
from datetime import datetime
from PIL import Image
import io
import reportlab
import PIL



# --- Third-Party Library Imports ---
import pandas as pd
import openpyxl  # Required for reading Excel headers
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
MIDTERM_TOTAL_MARKS=30

# Suppress a known warning from PyPDF2
warnings.filterwarnings("ignore", category=DeprecationWarning)
from PyPDF2 import PdfReader

try:
    from docx2pdf import convert
except ImportError:
    print("Warning: 'docx2pdf' module not found. PDF output will not be available.")
    convert = None

try:
    from endesive import pdf
    from cryptography.hazmat.primitives.serialization import load_pem_private_key
    from cryptography.x509 import load_pem_x509_certificate
except ImportError:
    print("Warning: Required crypto libraries 'endesive' and 'cryptography' not found.")
    print("PDF signing will not be available. Please run: pip install endesive cryptography")
    pdf = None

# ==============================================================================
# 0. CONSTANTS & CONFIGURATION
# ==============================================================================
MIDTERM_TOTAL_MARKS = 30
SEMESTER_MAPPING = {
    'i': 'I Year/ I semester',
    'ii': 'I Year/ II semester',
    'iii': 'II Year/ III semester',
    'iv': 'II Year/ IV semester',
    'v': 'III Year/ V semester',
    'vi': 'III Year/ VI semester',
    'vii': 'IV Year/ VII semester',
    'viii': 'IV Year/ VIII semester',
}

# ==============================================================================
# 1. PDF SIGNING UTILITY
# ==============================================================================
def sign_pdf(pdf_path, key_path, cert_path, image_path, password):
    """
    Signs a PDF with one digital signature and places
    a visible signature appearance on every page.
    """

    if not all([pdf_path, key_path, cert_path, image_path, password]):
        print("Skipping signing due to missing information.")
        return

    try:
        from datetime import datetime
        import io
        import traceback
        from PIL import Image
        from PyPDF2 import PdfReader, PdfWriter
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.utils import ImageReader
        from cryptography.hazmat.primitives.serialization import load_pem_private_key
        from cryptography.x509 import load_pem_x509_certificate
        from endesive import pdf

        # Signing timestamp
        date = datetime.now().strftime("D:%Y%m%d%H%M%S+05'30'")

        # Load private key and certificate
        with open(key_path, 'rb') as f:
            private_key = load_pem_private_key(f.read(), password=password.encode('utf-8'))
        with open(cert_path, 'rb') as f:
            certificate = load_pem_x509_certificate(f.read())

        # Read full original PDF once
        with open(pdf_path, 'rb') as f_in:
            pdf_data_initial = f_in.read()

        reader = PdfReader(pdf_path)
        writer = PdfWriter()
        page_count = len(reader.pages)

        # Load signature image
        with open(image_path, 'rb') as f_img:
            image_data = f_img.read()

        # Read image size
        with Image.open(io.BytesIO(image_data)) as img:
            img_width, img_height = img.size

        # Convert pixels to PDF points
        scale_factor = 0.75
        sig_width = img_width * scale_factor / 1.33
        sig_height = img_height * scale_factor / 1.33

        # Create signature box per page
        signatureboxes = []
        for page_index in range(page_count):
            page = reader.pages[page_index]
            page_width = float(page.mediabox.width)

            x2 = page_width - 50
            y2 = 50 + sig_height
            x1 = x2 - sig_width
            y1 = 50

            signatureboxes.append((page_index, x1, y1, x2, y2))

        # Draw signature appearance on PDF pages
        def add_signature_to_page(page, sig_coords):
            packet = io.BytesIO()
            can = canvas.Canvas(packet, pagesize=letter)

            x1, y1, x2, y2 = sig_coords
            sig_w = x2 - x1
            sig_h = y2 - y1

            signature_img = ImageReader(image_path)
            can.drawImage(signature_img, x1, y1, sig_w, sig_h, mask='auto')

            can.save()
            packet.seek(0)

            overlay_pdf = PdfReader(packet)
            overlay_page = overlay_pdf.pages[0]
            page.merge_page(overlay_page)
            return page

        # Apply visible signature image
        for (page_idx, x1, y1, x2, y2) in signatureboxes:
            page = reader.pages[page_idx]
            signed_page = add_signature_to_page(page, (x1, y1, x2, y2))
            writer.add_page(signed_page)

        # Save the visually signed PDF to a buffer
        buffer = io.BytesIO()
        writer.write(buffer)
        buffer.seek(0)
        updated_pdf_bytes = buffer.getvalue()

        # Prepare metadata for cryptographic signing
        signdata = {
            "sigflags": 3,
            "contact": "faculty.email@example.com",
            "location": "Manipal, India",
            "reason": "I am the author of this document",
            "signatureboxes": signatureboxes,
            "appearance": {
                "image": image_data,
                "fit": "stretch"
            },
            "signingdate": date
        }

        print(f"Signing PDF with one digital signature and {page_count} visible marks.")

        # Apply cryptographic signature
        # TEMP: Write the visually signed PDF ONLY
        with open("temp_visually_signed.pdf", "wb") as f:
            f.write(updated_pdf_bytes)
            print("Saved temp PDF without digital signature")
            return

            updated_pdf_bytes,
            signdata,
            key=private_key,
            cert=certificate,
            othercerts=[certificate]
        

        # Write signed output
        with open(pdf_path, "wb") as f_out:
            f_out.write(signed_pdf_bytes)

        print("Success. Signed file saved:", pdf_path)

    except Exception:
        print("Signing process failed.")
        traceback.print_exc()

# ==============================================================================
# 2. DATA READER
# ==============================================================================
class DataReader:
    """Reads data from an Excel or CSV file and returns it in a neutral format."""

    # --- UPDATED Column Mapping ---
    COLUMN_MAPPING = {
        # --- Mappings from your file ---
        'Roll Number': 'Register Number of the Student',
        'Student Name': 'Student Name',
        'Total (30) *': 'Midterm Exam Marks (Out of 30)',

        # --- Mappings for other potential columns ---
        'CGPA': 'CGPA (up to previous semester)',
        'Actions': 'Actions taken to improve performance',
        'Outcome': 'Outcome (Based on clearance in end-semester or makeup exam)',
        'Remarks': 'Remarks if any',
        'Sem': 'Semester',
        'Subject': 'Subject Name'
    }

    # --- UPDATED Subject Header Extractor ---
    def _extract_subject_from_header(self, file_path):
        """
        Attempts to read the subject name from cell B1 of the Excel file.
        """
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True, read_only=True)
            sheet = wb.active
            
            # Data is in cell B1, e.g., "Exam: ... / COMPUTER NETWORKS [CSE 3124]"
            subject_name_raw = sheet['B1'].value
            
            if subject_name_raw:
                # Get the last part after the final '/'
                subject_name = str(subject_name_raw).split('/')[-1].strip()
                # Optional: clean up brackets if they exist
                subject_name = subject_name.split('[')[0].strip()
                return subject_name
            return None
        except Exception as e:
            print(f"Warning: Could not auto-detect subject name from file header. Error: {e}")
            return None

    def read_data(self, file_path):
        """Reads the Excel file, returning both the data and the subject name."""
        subject_name = self._extract_subject_from_header(file_path)
        if not subject_name:
            subject_name = input("Could not auto-detect subject. Please enter Subject Name manually: ")

        try:
            # --- UPDATED skiprows to 1 ---
            df = pd.read_excel(file_path, skiprows=2)
            
            df.columns = df.columns.str.strip()
            df.rename(columns=self.COLUMN_MAPPING, inplace=True)

            required_cols = ['Register Number of the Student', 'Student Name', 'Midterm Exam Marks (Out of 30)']
            
            missing_cols = [col for col in required_cols if col not in df.columns]
            if missing_cols:
                print(f"\n--- CRITICAL ERROR ---")
                print(f"Error: The Excel file is missing required columns.")
                print(f"Missing: {missing_cols}")
                print(f"Please ensure your Excel headers match the keys in COLUMN_MAPPING.")
                print(f"Found headers: {list(df.columns)}")
                print("----------------------\n")
                return None, None

            reg_col = 'Register Number of the Student'
            df[reg_col] = df[reg_col].astype(str).str.replace(r'\.0$', '', regex=True)

            return df.to_dict('records'), subject_name
        except FileNotFoundError:
            print(f"Error: The file '{file_path}' was not found.")
            return None, None
        except Exception as e:
            print(f"An error occurred while reading the data file: {e}")
            traceback.print_exc()
            return None, None

# ==============================================================================
# 3. DATA PROCESSOR
# ==============================================================================
class DocxWriter:
    """Takes a Document object and saves it to a .docx file."""
    def write(self, doc, output_filename, **kwargs):
        try:
            doc.save(output_filename)
            print(f"\nSuccess! Report generated as '{output_filename}' ✨")
        except Exception as e:
            print(f"\nError: Could not save the file. Details: {e}")

class PdfWriter:
    """Creates a DOCX, converts it to PDF, and optionally signs it."""
    def write(self, doc, output_filename, sign_info=None, format_choice=None):
        if convert is None:
            print("\nError: 'docx2pdf' not installed. Cannot generate PDF.")
            return

        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_docx = os.path.join(temp_dir, "temp_report.docx")
                doc.save(temp_docx)
                convert(temp_docx, output_filename)

            time.sleep(2)

            if platform.system() == "Darwin":
                time.sleep(1)

            print(f"\nSuccess! PDF generated as '{output_filename}' ✨")

            if sign_info and sign_info.get('should_sign'):
                if format_choice in ['1', '2', '4', '5']:
                    print("Proceeding to sign the PDF...")
                    sign_pdf(
                        pdf_path=output_filename,
                        key_path=sign_info['key_path'],
                        cert_path=sign_info['cert_path'],
                        image_path=sign_info['image_path'],
                        password=sign_info['password']
                    )
                else:
                    print("Skipping signature for Format 3 (Summary Report).")

        except Exception as e:
            print(f"\nError: Could not save or convert the file. Details: {e}")
            print("Please ensure Microsoft Word (on Windows) or LibreOffice (on macOS/Linux) is installed and accessible.")

# ==============================================================================
# 4. REPORT FORMATTERS
# ==============================================================================
class BaseFormatter:
    """Base class for all formatters with shared helper methods."""
    COMIC_SANS = "Brush Script MT Italic"
    
    # --- ADDED Font Name ---
    FONT_NAME = "Times New Roman" 

    def get_year_semester_string(self, roman_numeral):
        return SEMESTER_MAPPING.get(str(roman_numeral).strip().lower(), str(roman_numeral))

    def set_cell_properties(self, cell, text, bold=False, font_size=10, align='LEFT', valign='TOP', font_name=None):
        cell.text = ''
        p = cell.add_paragraph()
        run = p.add_run(str(text))
        run.font.size = Pt(font_size)
        run.bold = bold
        if font_name:
            run.font.name = font_name
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, str(align).upper(), WD_ALIGN_PARAGRAPH.LEFT)
        cell.vertical_alignment = getattr(WD_ALIGN_VERTICAL, str(valign).upper(), WD_ALIGN_VERTICAL.TOP)

    def add_signature_line(self, doc_or_cell):
        p = doc_or_cell.add_paragraph()
        p.add_run("\n\n" + "_" * 40 + "\n")
        p.add_run("Signature of the\nsubject teacher / class coordinator")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def _add_document_header(self, cell):
        p1 = cell.add_paragraph(); p1.add_run('Manipal Institute of Technology').bold = True; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2 = cell.add_paragraph(); p2.add_run('MAHE Manipal').bold = True; p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3 = cell.add_paragraph(); p3.add_run('Computer Science and Engineering Department').bold = True; p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _create_format1_content(self, doc, student, slow_threshold, fast_threshold):
        doc.add_heading('Format 1. Assessment of the learning levels of the students:', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
        container_table = doc.add_table(rows=5, cols=1)
        container_table.style = 'Table Grid'
        self._add_document_header(container_table.cell(0, 0))
        
        student_info_table = container_table.cell(1, 0).add_table(rows=4, cols=2)
        self.set_cell_properties(student_info_table.cell(0, 0), 'Name of the Student:')
        self.set_cell_properties(student_info_table.cell(0, 1), student.get('Student Name', ''), font_name=self.FONT_NAME)
        self.set_cell_properties(student_info_table.cell(1, 0), 'Registration Number:')
        self.set_cell_properties(student_info_table.cell(1, 1), student.get('Register Number of the Student', ''), font_name=self.FONT_NAME)
        self.set_cell_properties(student_info_table.cell(2, 0), 'Course:')
        self.set_cell_properties(student_info_table.cell(2, 1), str(student.get('Subject Name', '')).title(), font_name=self.FONT_NAME)
        self.set_cell_properties(student_info_table.cell(3, 0), 'Year /semester:')
        self.set_cell_properties(student_info_table.cell(3, 1), self.get_year_semester_string(student.get('Semester', '')), font_name=self.FONT_NAME)
        
        params_table = container_table.cell(2, 0).add_table(rows=3, cols=4)
        params_table.style = 'Table Grid'
        params_table.cell(0, 2).merge(params_table.cell(0, 3))
        self.set_cell_properties(params_table.cell(0, 0), 'Sr. No.', bold=True, align='CENTER')
        self.set_cell_properties(params_table.cell(0, 1), 'Parameter', bold=True, align='CENTER')
        self.set_cell_properties(params_table.cell(0, 2), 'Weightage in Percentage', bold=True, align='CENTER')
        self.set_cell_properties(params_table.cell(1, 0), '1', align='CENTER')
        self.set_cell_properties(params_table.cell(1, 1), f"Scores obtained by student class test / internal examination...\nConsidered Midterm exam conducted for {MIDTERM_TOTAL_MARKS}M:")
        self.set_cell_properties(params_table.cell(1, 2), f"{student.get('MidtermPercentage', 0):.2f}", align='CENTER', font_name=self.FONT_NAME)
        self.set_cell_properties(params_table.cell(1, 3), "> %", align='CENTER')
        self.set_cell_properties(params_table.cell(2, 0), '2', align='CENTER')
        self.set_cell_properties(params_table.cell(2, 1), 'Performance of students in preceding university examination')
        self.set_cell_properties(params_table.cell(2, 2), str(student.get('CGPA (up to previous semester)', '')), align='CENTER', font_name=self.FONT_NAME)
        self.set_cell_properties(params_table.cell(2, 3), "> %", align='CENTER')
        
        container_table.cell(3, 0).text = "Total Weightage"
        footer_cell = container_table.cell(4, 0)
        footer_cell.add_paragraph(f"1. Midterm score less than {slow_threshold}% considered as a slow learner")
        footer_cell.add_paragraph(f"2. Midterm score more than {fast_threshold}% considered as an advanced learner **")
        p_date = footer_cell.add_paragraph()
        p_date.add_run(f"Date: {datetime.now().strftime('%d-%m-%Y')}").font.name = self.FONT_NAME
        self.add_signature_line(footer_cell)

    def _create_format2_content(self, doc, student):
        doc.add_heading('Format -2 Report of performance/ improvement for slow and advanced learners', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_table = doc.add_table(rows=1, cols=1)
        self._add_document_header(header_table.cell(0,0))
        
        content_table = doc.add_table(rows=8, cols=2)
        content_table.style = 'Table Grid'
        self.set_cell_properties(content_table.cell(0, 0), '1. Registration Number')
        self.set_cell_properties(content_table.cell(0, 1), student.get('Register Number of the Student', ''), font_name=self.FONT_NAME)
        self.set_cell_properties(content_table.cell(1, 0), '2. Name of the student')
        self.set_cell_properties(content_table.cell(1, 1), student.get('Student Name', ''), font_name=self.FONT_NAME)
        self.set_cell_properties(content_table.cell(2, 0), '3. Course')
        self.set_cell_properties(content_table.cell(2, 1), str(student.get('Subject Name', '')).title(), font_name=self.FONT_NAME)
        self.set_cell_properties(content_table.cell(3, 0), '4. Year/Semester')
        self.set_cell_properties(content_table.cell(3, 1), self.get_year_semester_string(student.get('Semester', '')), font_name=self.FONT_NAME)
        self.set_cell_properties(content_table.cell(4, 0), '5. Midterm Percentage')
        self.set_cell_properties(content_table.cell(4, 1), f"{student.get('MidtermPercentage', 0):.2f}%", font_name=self.FONT_NAME)
        self.set_cell_properties(content_table.cell(5, 0), '6. Activities/ Measure/special programs\ntaken to improve the performance')
        self.set_cell_properties(content_table.cell(5, 1), str(student.get('Actions taken to improve performance', '')).replace(';', '\n'), font_name=self.FONT_NAME)
        self.set_cell_properties(content_table.cell(6, 0), '7. Progress')
        self.set_cell_properties(content_table.cell(6, 1), str(student.get('Outcome (Based on clearance in end-semester or makeup exam)', '')), font_name=self.FONT_NAME)
        self.set_cell_properties(content_table.cell(7, 0), 'Comments/remarks')
        self.set_cell_properties(content_table.cell(7, 1), str(student.get('Remarks if any', '')), font_name=self.FONT_NAME)
        
        p_date = doc.add_paragraph()
        p_date.add_run(f"\nDate:{datetime.now().strftime('%d-%m-%Y')}").font.name = self.FONT_NAME
        self.add_signature_line(doc)

    def _generate_pages(self, doc, students, content_method, *args):
        for i, student in enumerate(students):
            content_method(doc, student, *args)
            if i < len(students) - 1:
                doc.add_page_break()
        return doc

class Format1DocxFormatter(BaseFormatter):
    def format(self, students, slow_threshold, fast_threshold):
        doc = Document()
        return self._generate_pages(doc, students, self._create_format1_content, slow_threshold, fast_threshold)

class Format2DocxFormatter(BaseFormatter):
    def format(self, students, slow_threshold, fast_threshold):
        doc = Document()
        return self._generate_pages(doc, students, self._create_format2_content)

class Format3DocxFormatter(BaseFormatter):
    def format(self, students, slow_threshold, fast_threshold):
        doc = Document()
        df = pd.DataFrame(students)
        grouped = df.groupby(['Subject Name', 'Semester'])
        for i, ((subject, semester), group) in enumerate(grouped):
            doc.add_paragraph(f"Course: {str(subject).title()}", style='Heading 3')
            doc.add_paragraph(f"Year /Semester: {self.get_year_semester_string(semester)}", style='Heading 3')
            summary_cols = ['Sl. No', 'Reg Number', 'Name of the student', 'Midterm Percentage', 'Progress']
            table = doc.add_table(rows=1, cols=len(summary_cols))
            table.style = 'Table Grid'
            for j, col_name in enumerate(summary_cols):
                self.set_cell_properties(table.cell(0, j), col_name, bold=True)
            for index, row_data in group.reset_index(drop=True).iterrows():
                row_cells = table.add_row().cells
                self.set_cell_properties(row_cells[0], str(index + 1), font_name=self.FONT_NAME)
                self.set_cell_properties(row_cells[1], str(row_data.get('Register Number of the Student', '')), font_name=self.FONT_NAME)
                self.set_cell_properties(row_cells[2], str(row_data.get('Student Name', '')), font_name=self.FONT_NAME)
                self.set_cell_properties(row_cells[3], f"{row_data.get('MidtermPercentage', 0):.2f}", font_name=self.FONT_NAME)
                self.set_cell_properties(row_cells[4], str(row_data.get('Outcome (Based on clearance in end-semester or makeup exam)', '')), font_name=self.FONT_NAME)
            if i < len(grouped) - 1:
                doc.add_page_break()
        return doc

class Format1And2DocxFormatter(BaseFormatter):
    def format(self, students, slow_threshold, fast_threshold):
        doc = Document()
        for i, student in enumerate(students):
            self._create_format1_content(doc, student, slow_threshold, fast_threshold)
            doc.add_page_break()
            self._create_format2_content(doc, student)
            if i < len(students) - 1:
                doc.add_page_break()
        return doc

# ==============================================================================
# 5. FACTORIES
# ==============================================================================
def get_writer(output_type):
    if output_type == 'word': return DocxWriter()
    if output_type == 'pdf': return PdfWriter()
    raise ValueError(f"Unsupported output type: {output_type}")

def get_formatter(format_choice):
    formatters = {
        '1': Format1DocxFormatter(), '2': Format2DocxFormatter(),
        '3': Format3DocxFormatter(), '4': Format1And2DocxFormatter(),
    }
    formatter = formatters.get(format_choice)
    if formatter: return formatter
    raise ValueError(f"Unsupported format choice '{format_choice}'")

# ==============================================================================
# 6. DATA PROCESSOR
# ==============================================================================
class StudentDataProcessor:
    def _calculate_midterm_percentage(self, marks):
        try:
            return (float(marks) / MIDTERM_TOTAL_MARKS) * 100
        except (ValueError, TypeError):
            return 0
            
    # --- UPDATED process_data method ---
    def process_data(self, all_student_data, common_comment, subject_name, semester):
        """
        Processes student data, calculates percentages, and adds common actions.
        """
        learner_action_key = 'Actions taken to improve performance'
        
        for student in all_student_data:
            student['MidtermPercentage'] = self._calculate_midterm_percentage(
                student.get('Midterm Exam Marks (Out of 30)')
            )
            
            # --- FIX: Manually inject subject and semester ---
            student['Subject Name'] = subject_name
            student['Semester'] = semester
            
            # --- Logic to add the common comment ---
            if common_comment:
                existing_actions = str(student.get(learner_action_key, '')).strip()
                if existing_actions:
                    # Append if actions already exist
                    student[learner_action_key] = f"{existing_actions}; {common_comment}"
                else:
                    # Set if no actions exist
                    student[learner_action_key] = common_comment
                    
        return all_student_data
        
    def filter_students(self, students, semester, subject, learner_type, slow_thresh, fast_thresh):
        # Filter by course/semester is now redundant since it's injected,
        # but we keep it for potential future use cases.
        filtered_by_course = [
            s for s in students if
            (semester == 'all' or s['Semester'] == semester) and
            (subject == 'all' or s['Subject Name'] == subject)
        ]
        
        if learner_type == 'slow':
            final_filtered = [s for s in filtered_by_course if s['MidtermPercentage'] <= slow_thresh]
        else:
            final_filtered = [s for s in filtered_by_course if s['MidtermPercentage'] >= fast_thresh]
            
        if subject == 'all':
            final_filtered.sort(key=lambda s: (s.get('Subject Name', ''), s.get('Register Number of the Student', '')))
        else:
            final_filtered.sort(key=lambda s: s.get('Register Number of the Student', ''))
            
        return final_filtered

# ==============================================================================
# 7. CONTROLLER
# ==============================================================================
class ReportController:
    def __init__(self, excel_path, format_choice, learner_type, slow_thresh, fast_thresh, output_type, semester, sign_info, common_comment):
        self.excel_path = excel_path
        self.format_choice = format_choice
        self.learner_type = learner_type
        self.slow_threshold = slow_thresh
        self.fast_threshold = fast_thresh
        self.output_type = output_type
        self.subject = ""
        self.semester = semester.lower().strip()
        self.sign_info = sign_info
        self.common_comment = common_comment
        self.reader = DataReader()
        self.processor = StudentDataProcessor()
        self.writer = get_writer(output_type)

    def run(self):
        all_student_data, detected_subject = self.reader.read_data(self.excel_path)
        if not all_student_data:
            return
        
        # Use auto-detected subject. self.semester is from user input
        self.subject = detected_subject.lower().strip()

        # --- UPDATED call to process_data ---
        processed_students = self.processor.process_data(
            all_student_data, self.common_comment, self.subject, self.semester
        )
        
        # Now we filter based on the data we just injected
        final_filtered_students = self.processor.filter_students(
            processed_students,
            self.semester,  # Use the semester from user input
            self.subject,   # Use the subject from the file header
            self.learner_type,
            self.slow_threshold,
            self.fast_threshold
        )
        
        if not final_filtered_students:
            print(f"\nNo students found for the selected criteria.")
            print(f"(Semester: '{self.semester}', Subject: '{self.subject}', Type: '{self.learner_type}')")
            return

        print(f"\nFound {len(final_filtered_students)} {self.learner_type} learners.")
        
        # --- Create dynamic output directory structure ---
        date_str = datetime.now().strftime('%d_%m_%y')
        base_dir = "Learner Monitor Reports"
        learner_folder = f"{self.learner_type.title()} Learners"
        
        semester_name_for_file = self.semester.upper()
        subject_name_for_file = self.subject.replace(' ', '_').title()
        subject_name_for_file = re.sub(r'[\\/*?:"<>|]', "", subject_name_for_file) # Remove illegal chars

        output_dir = os.path.join(base_dir, learner_folder, f"Semester_{semester_name_for_file}", subject_name_for_file)
        os.makedirs(output_dir, exist_ok=True)
        
        if self.format_choice == '5':
            self._generate_all_formats(final_filtered_students, output_dir, date_str, semester_name_for_file, subject_name_for_file)
            return

        formatter = get_formatter(self.format_choice)
        output_object = formatter.format(final_filtered_students, self.slow_threshold, self.fast_threshold)
        
        file_extension = 'docx' if self.output_type == 'word' else 'pdf'
        report_name = {'1': 'Format1', '2': 'Format2', '3': 'Summary', '4': 'Combined'}.get(self.format_choice, "Report")
        output_filename = f'{subject_name_for_file}_{semester_name_for_file}_{self.learner_type.title()}Learner_{report_name}_{date_str}.{file_extension}'
        full_output_path = os.path.join(output_dir, output_filename)
        
        self.writer.write(output_object, full_output_path, sign_info=self.sign_info, format_choice=self.format_choice)

    def _generate_all_formats(self, students, output_dir, date_str, semester_name_for_file, subject_name_for_file):
        file_extension = 'docx' if self.output_type == 'word' else 'pdf'
        
        # Combined Report (Format 1 & 2)
        combined_filename = f'{subject_name_for_file}_{semester_name_for_file}_{self.learner_type.title()}Learner_Combined_Report_{date_str}.{file_extension}'
        f1_and_2_formatter = Format1And2DocxFormatter()
        doc1 = f1_and_2_formatter.format(students, self.slow_threshold, self.fast_threshold)
        self.writer.write(doc1, os.path.join(output_dir, combined_filename), sign_info=self.sign_info, format_choice='4') # Sign as if format 4

        # Summary Report (Format 3)
        summary_filename = f'{subject_name_for_file}_{semester_name_for_file}_{self.learner_type.title()}Learner_Summary_Report_{date_str}.{file_extension}'
        f3_formatter = Format3DocxFormatter()
        doc2 = f3_formatter.format(students, self.slow_threshold, self.fast_threshold)
        self.writer.write(doc2, os.path.join(output_dir, summary_filename), sign_info=self.sign_info, format_choice='3') # Don't sign summary

# ==============================================================================
# 8. MAIN EXECUTION BLOCK (User Interface)
# ==============================================================================
def get_valid_input(prompt, valid_options=None, input_type=str):
    
    while True:
        user_input = input(prompt).strip()
        if not user_input and input_type is not str:
             print("This field cannot be empty.")
             continue
        if not user_input and input_type is str:
             return user_input # Allow empty string for manual subject input
        try:
            converted_input = input_type(user_input)
            if valid_options and str(converted_input) not in valid_options:
                print(f"Invalid input. Please choose from {valid_options}.")
                continue
            return converted_input
        except ValueError:
            print("Invalid input. Please enter a valid number.")

if __name__ == "__main__":
    print("-" * 50)
    print("Welcome to the Student Learner Report Generator.")
    print("-" * 50)

    excel_file = input("Enter the Excel file name or full file path: ")
    if not os.path.exists(excel_file):
        print(f"Error: The file '{excel_file}' does not exist.")
        exit()
    
    semester_filter = get_valid_input("Enter Semester (e.g., 'III', 'V', etc.): ", input_type=str)
    if not semester_filter:
        print("Error: Semester cannot be empty.")
        exit()

    output_format = get_valid_input("Choose output format ('word' or 'pdf'): ", ['word', 'pdf'])
    learner = get_valid_input("Generate report for 'fast' or 'slow' learners? ", ['fast', 'slow'])
    
    if learner == 'slow':
        comment = input("Enter common action for SLOW learners (e.g., 'Remedial classes conducted'): ")
    else:
        comment = input("Enter common action for FAST learners (e.g., 'Advanced assignments given'): ")

    slow_thresh = get_valid_input("Enter percentage threshold for SLOW learners (e.g., 40): ", input_type=float)
    fast_thresh = get_valid_input("Enter percentage threshold for FAST learners (e.g., 80): ", input_type=float)
    
    print("\nPlease choose a report format:")
    print(" 1: Format 1 - Assessment of learning levels (One page per student)")
    print(" 2: Format 2 - Report of performance/improvement (One page per student)")
    print(" 3: Format 3 - Tabular Summary Report (All students on one page)")
    print(" 4: Combined Format 1 & 2 (Two pages per student in one file)")
    print(" 5: All Formats (Generates two separate files: Combined and Summary)")
    format_num = get_valid_input("Enter your choice (1, 2, 3, 4, or 5): ", ['1', '2', '3', '4', '5'])
    
    signing_data = {'should_sign': False}
    if output_format == 'pdf':
        sign_choice = get_valid_input("\nDo you want to digitally sign the PDF report? (y/n): ", ['y', 'n'])
        if sign_choice.lower() == 'y':
            if pdf is None:
                print("\nCannot sign PDF. Please install required libraries with: pip install endesive cryptography")
            else:
                print("\nPlease provide paths to your signing assets.")
                key_file = input("Path to your private key file (.pem): ")
                cert_file = input("Path to your certificate file (.pem): ")
                image_file = input("Path to your signature image file (.png): ")
                password = getpass.getpass("Enter the password for your private key: ")
                signing_data = {'should_sign': True, 'key_path': key_file, 'cert_path': cert_file, 'image_path': image_file, 'password': password}
    
    controller = ReportController(excel_file, format_num, learner, slow_thresh, fast_thresh, output_format, semester_filter, signing_data, comment)
    controller.run()
    
    print("\nReport generation process finished.")
    input("Press Enter to exit.")