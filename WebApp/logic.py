# ==============================================================================
# REPORT GENERATION LOGIC LIBRARY - FINAL PRODUCTION VERSION
# ==============================================================================

import os
import tempfile
import platform
import re
import traceback
import warnings
import subprocess
from datetime import datetime

import pandas as pd
from endesive import pdf
from cryptography.hazmat.primitives.serialization import load_pem_private_key
from cryptography.x509 import load_pem_x509_certificate

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

warnings.filterwarnings("ignore", category=DeprecationWarning)

# --- CONFIGURATION ---
MIDTERM_TOTAL_MARKS = 30
SEMESTER_MAPPING = {
    'i': 'I Year/ I semester', 'ii': 'I Year/ II semester', 'iii': 'II Year/ III semester',
    'iv': 'II Year/ IV semester', 'v': 'III Year/ V semester', 'vi': 'III Year/ VI semester',
    'vii': 'IV Year/ VII semester', 'viii': 'IV Year/ VIII semester',
}

def get_libreoffice_command():
    system = platform.system()
    if system == 'Darwin': return '/Applications/LibreOffice.app/Contents/MacOS/soffice'
    elif system == 'Windows': return r'C:\Program Files\LibreOffice\program\soffice.exe'
    else: return 'libreoffice'

def sign_pdf(pdf_path, key_path, cert_path, image_path, password):
    try:
        date = datetime.now().strftime('D:%Y%m%d%H%M%S+05\'30\'')
        with open(key_path, 'rb') as f: 
            private_key = load_pem_private_key(f.read(), password=password.encode('utf-8') if password else None)
        with open(cert_path, 'rb') as f: 
            certificate = load_pem_x509_certificate(f.read())
        with open(pdf_path, 'rb') as f_in: 
            pdf_data = f_in.read()

        signdata = {
            'sigflags': 3, 'contact': 'faculty@example.com', 'location': 'Manipal',
            'reason': 'Verified Report', 'signingdate': date, 'page': 0
        }
        signed_pdf_bytes = pdf.cms.sign(pdf_data, signdata, key=private_key, cert=certificate, othercerts=())
        with open(pdf_path, 'wb') as f_out: 
            f_out.write(pdf_data + signed_pdf_bytes)
        return True
    except Exception:
        return False

def normalize_registration_number(reg_num):
    if pd.isna(reg_num) or reg_num is None:
        return ''
    normalized = str(reg_num).strip()
    normalized = normalized.replace('.0', '').replace(' ', '').upper()
    return normalized

# --- DATA READER ---
class DataReader:
    COLUMN_MAPPING = {
        'Roll Number': 'Register Number of the Student', 
        'Student Name': 'Student Name',
        'Total (30) *': 'Midterm Exam Marks (Out of 30)', 
        'Student Viewed': 'Did student view the paper'
    }

    def _extract_subject_from_header(self, file_path):
        try:
            engine = 'xlrd' if file_path.lower().endswith('.xls') else 'openpyxl'
            df_header = pd.read_excel(file_path, engine=engine, nrows=5, header=None)
            for val in df_header.iloc[:, 0]:
                if val and isinstance(val, str) and "Exam:" in val:
                    last_slash = val.rfind('/')
                    first_bracket = val.find('[')
                    last_bracket = val.find(']')
                    if last_slash != -1 and first_bracket != -1:
                        name = val[last_slash + 1 : first_bracket].strip()
                        code = val[first_bracket + 1 : last_bracket].strip() if last_bracket != -1 else ""
                        return f"{name} ({code})" if code else name
            return None
        except Exception: return None

    def read_data(self, file_path):
        subject_name = self._extract_subject_from_header(file_path)
        if not subject_name: raise ValueError("Could not auto-detect subject name.")
        try:
            engine = 'xlrd' if file_path.lower().endswith('.xls') else 'openpyxl'
            df = pd.read_excel(file_path, skiprows=2, engine=engine)
            df.columns = df.columns.str.strip()
            df.rename(columns=self.COLUMN_MAPPING, inplace=True)
            reg_col = 'Register Number of the Student'
            if reg_col in df.columns:
                df[reg_col] = df[reg_col].apply(normalize_registration_number)
            return df.to_dict('records'), subject_name
        except Exception: raise

    def read_cgpa_map(self, file_path):
        if not file_path or not os.path.exists(file_path):
            return {}
        try:
            engine = 'xlrd' if file_path.lower().endswith('.xls') else 'openpyxl'
            if file_path.lower().endswith('.csv'): 
                df = pd.read_csv(file_path)
            else:
                xl_file = pd.ExcelFile(file_path)
                best_sheet = xl_file.sheet_names[0]
                if len(xl_file.sheet_names) > 1:
                    max_rows = 0
                    for sheet in xl_file.sheet_names:
                        temp_df = pd.read_excel(file_path, sheet_name=sheet, engine=engine)
                        if len(temp_df) > max_rows:
                            max_rows = len(temp_df)
                            best_sheet = sheet
                df = pd.read_excel(file_path, sheet_name=best_sheet, engine=engine)

            if len(df.columns) >= 2:
                # Col 1 (Index 0) = Reg No, Col 2 (Index 1) = CGPA
                roll_col, cgpa_col = df.columns[0], df.columns[1]
                df[roll_col] = df[roll_col].apply(normalize_registration_number)
                return pd.Series(df[cgpa_col].values, index=df[roll_col]).to_dict()
            return {}
        except Exception: return {}

    def read_grade_map(self, file_path, course_code=None):
        if not file_path or not os.path.exists(file_path):
            return {}
        try:
            engine = 'xlrd' if file_path.lower().endswith('.xls') else 'openpyxl'
            df = pd.read_csv(file_path) if file_path.lower().endswith('.csv') else pd.read_excel(file_path, engine=engine)
            
            # Index 0: Enrollment, Index 1: Course Code, Index 2: Grade
            if len(df.columns) >= 3:
                enroll_col, course_col, grade_col = df.columns[0], df.columns[1], df.columns[2]
                
                if course_code:
                    target = course_code.split('(')[-1].replace(')', '').strip() if '(' in course_code else course_code
                    df = df[df[course_col].astype(str).str.contains(target, case=False, na=False)]
                
                df[enroll_col] = df[enroll_col].apply(normalize_registration_number)
                return pd.Series(df[grade_col].values, index=df[enroll_col]).to_dict()
            return {}
        except Exception: return {}

# --- DATA PROCESSOR ---
class StudentDataProcessor:
    def _calculate_midterm_percentage(self, marks):
        try: return (float(marks) / MIDTERM_TOTAL_MARKS) * 100
        except: return 0

    def process_data(self, all_student_data, subject_name, semester, common_comment, cgpa_map=None, grade_map=None):
        cgpa_map = cgpa_map or {}
        grade_map = grade_map or {} 
        for student in all_student_data:
            student['MidtermPercentage'] = self._calculate_midterm_percentage(student.get('Midterm Exam Marks (Out of 30)'))
            student['Subject Name'] = str(subject_name).strip()
            student['Semester'] = str(semester).strip().lower()
            roll_no = normalize_registration_number(student.get('Register Number of the Student', ''))
            student['CGPA (up to previous semester)'] = cgpa_map.get(roll_no, '')
            student['Actions taken to improve performance'] = common_comment
            
            grade = str(grade_map.get(roll_no, '')).strip().upper()
            if not grade or grade in ['NAN', 'NONE']:
                student['Outcome (Based on clearance in end-semester or makeup exam)'] = ''
            elif grade in {'A+', 'A', 'B', 'C', 'D', 'E', 'S'}:
                student['Outcome (Based on clearance in end-semester or makeup exam)'] = 'Improved'
            else:
                student['Outcome (Based on clearance in end-semester or makeup exam)'] = 'Not Improved'
        return all_student_data

    def filter_students(self, students, learner_type, slow_thresh, advanced_thresh):
        if learner_type == 'slow':
            filtered = [s for s in students if s['MidtermPercentage'] < slow_thresh]
        else:
            filtered = [s for s in students if s['MidtermPercentage'] > advanced_thresh]
        filtered.sort(key=lambda s: s.get('Register Number of the Student', ''))
        return filtered

# --- FORMATTERS ---
class BaseFormatter:
    BODY_FONT = "Times New Roman"
    def __init__(self):
        self.signature_image_path = None
        self.faculty_name = None
        self.learner_type = None

    def get_year_semester_string(self, roman_numeral): 
        return SEMESTER_MAPPING.get(str(roman_numeral).strip().lower(), str(roman_numeral))
    
    def set_cell_properties(self, cell, text, bold=False, font_size=10, align='LEFT', valign='TOP', font_name=None):
        cell.text = ''
        p = cell.add_paragraph()
        run = p.add_run(str(text))
        run.font.size = Pt(font_size); run.bold = bold
        if font_name: run.font.name = font_name
        p.alignment = getattr(WD_ALIGN_PARAGRAPH, str(align).upper(), WD_ALIGN_PARAGRAPH.LEFT)
        cell.vertical_alignment = getattr(WD_ALIGN_VERTICAL, str(valign).upper(), WD_ALIGN_VERTICAL.TOP)
    
    def add_signature_line(self, doc_or_cell):
        p = doc_or_cell.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if self.signature_image_path and os.path.exists(self.signature_image_path):
            try:
                run = p.add_run(); run.add_picture(self.signature_image_path, width=Inches(0.8)); run.add_break()
            except: pass
        p.add_run("_" * 40 + "\n")
        if self.faculty_name: p.add_run(f"{self.faculty_name}\n")
        p.add_run("Signature of the\nsubject teacher / class coordinator")

    def _add_document_header(self, cell):
        for line in ['Manipal Institute of Technology', 'MAHE Manipal', 'Computer Science and Engineering Department']:
            p = cell.add_paragraph(); p.add_run(line).bold = True; p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _create_format1_content(self, doc, student, slow_threshold, fast_threshold):
        doc.add_heading('Format 1. Assessment of the learning levels of the students:', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
        ct = doc.add_table(rows=5, cols=1); ct.style = 'Table Grid'; self._add_document_header(ct.cell(0, 0))
        st_table = ct.cell(1, 0).add_table(rows=4, cols=2)
        
        mapping = [('Name of the Student:', 'Student Name'), ('Registration Number:', 'Register Number of the Student'), 
                   ('Course:', 'Subject Name'), ('Year /semester:', 'Semester')]
        for i, (label, key) in enumerate(mapping):
            self.set_cell_properties(st_table.cell(i, 0), label)
            val = student.get(key, '') if key != 'Semester' else self.get_year_semester_string(student.get(key, ''))
            self.set_cell_properties(st_table.cell(i, 1), str(val).upper() if key == 'Subject Name' else val, font_name=self.BODY_FONT)

        pt = ct.cell(2, 0).add_table(rows=3, cols=4); pt.style = 'Table Grid'; pt.cell(0, 2).merge(pt.cell(0, 3))
        headers = [('Sr. No.', 0), ('Parameter', 1), ('Weightage in Percentage', 2)]
        for text, col in headers: self.set_cell_properties(pt.cell(0, col), text, bold=True, align='CENTER')
        
        self.set_cell_properties(pt.cell(1, 0), '1', align='CENTER')
        self.set_cell_properties(pt.cell(1, 1), f"Scores obtained by student class test / internal examination...\nConsidered Midterm exam conducted for {MIDTERM_TOTAL_MARKS}M:")
        self.set_cell_properties(pt.cell(1, 2), f"{student.get('MidtermPercentage', 0):.2f}", align='CENTER', font_name=self.BODY_FONT)
        self.set_cell_properties(pt.cell(1, 3), "> %", align='CENTER')
        
        self.set_cell_properties(pt.cell(2, 0), '2', align='CENTER')
        self.set_cell_properties(pt.cell(2, 1), 'Performance of students in preceding university examination')
        self.set_cell_properties(pt.cell(2, 2), str(student.get('CGPA (up to previous semester)', '')), align='CENTER', font_name=self.BODY_FONT)
        self.set_cell_properties(pt.cell(2, 3), "> %", align='CENTER')

        ct.cell(3, 0).text = "Total Weightage"; fc = ct.cell(4, 0)
        p1 = fc.add_paragraph(); p1.add_run(f"1. Midterm score less than {slow_threshold}% considered as a ")
        r1 = p1.add_run("slow learner"); (self.learner_type == 'slow') and (setattr(r1.font, 'underline', True), setattr(r1.font.color, 'rgb', RGBColor(255, 0, 0)))
        
        p2 = fc.add_paragraph(); p2.add_run(f"2. Midterm score more than {fast_threshold}% considered as an ")
        r2 = p2.add_run("advanced learner"); (self.learner_type == 'advanced') and (setattr(r2.font, 'underline', True), setattr(r2.font.color, 'rgb', RGBColor(255, 0, 0)))
        
        p2.add_run(" **"); pd_ = fc.add_paragraph(); pd_.add_run(f"Date: {datetime.now().strftime('%d-%m-%Y')}").font.name = self.BODY_FONT; self.add_signature_line(fc)

    def _create_format2_content(self, doc, student):
        h = doc.add_paragraph(); h.style = 'Heading 2'; h.alignment = WD_ALIGN_PARAGRAPH.CENTER; h.add_run('Format -2 Report of performance/ improvement for ')
        r1 = h.add_run('slow'); (self.learner_type == 'slow') and (setattr(r1.font, 'underline', True), setattr(r1.font.color, 'rgb', RGBColor(255, 0, 0)))
        h.add_run(' and '); r2 = h.add_run('advanced'); (self.learner_type == 'advanced') and (setattr(r2.font, 'underline', True), setattr(r2.font.color, 'rgb', RGBColor(255, 0, 0)))
        h.add_run(' learners')
        
        ht = doc.add_table(rows=1, cols=1); self._add_document_header(ht.cell(0,0))
        ct = doc.add_table(rows=8, cols=2); ct.style = 'Table Grid'
        fields = [('1. Registration Number', 'Register Number of the Student'), ('2. Name of the student', 'Student Name'), 
                  ('3. Course', 'Subject Name'), ('4. Year/Semester', 'Semester'), ('5. Midterm Percentage', 'MidtermPercentage'),
                  ('6. Activities/ Measure/special programs\ntaken to improve the performance', 'Actions taken to improve performance'),
                  ('7. Progress', 'Outcome (Based on clearance in end-semester or makeup exam)'), ('Comments/remarks', 'Remarks if any')]
        
        for i, (label, key) in enumerate(fields):
            self.set_cell_properties(ct.cell(i, 0), label)
            val = student.get(key, '')
            if key == 'Semester': val = self.get_year_semester_string(val)
            elif key == 'MidtermPercentage': val = f"{val:.2f}%"
            elif key == 'Subject Name': val = str(val).upper()
            self.set_cell_properties(ct.cell(i, 1), str(val).replace(';', '\n'), font_name=self.BODY_FONT)
            
        pd_ = doc.add_paragraph(); pd_.add_run(f"\nDate:{datetime.now().strftime('%d-%m-%Y')}").font.name = self.BODY_FONT; self.add_signature_line(doc)

class Format1DocxFormatter(BaseFormatter):
    def format(self, s, st, ft): 
        doc = Document(); [setattr(sec, 'top_margin', Inches(0.5)) for sec in doc.sections]
        for i, student in enumerate(s): self._create_format1_content(doc, student, st, ft); (i < len(s)-1) and doc.add_page_break()
        return doc

class Format2DocxFormatter(BaseFormatter):
    def format(self, s, st, ft): 
        doc = Document(); [setattr(sec, 'top_margin', Inches(0.5)) for sec in doc.sections]
        for i, student in enumerate(s): self._create_format2_content(doc, student); (i < len(s)-1) and doc.add_page_break()
        return doc

class Format3DocxFormatter(BaseFormatter):
    def format(self, students, st, ft):
        doc = Document(); [setattr(sec, 'top_margin', Inches(0.5)) for sec in doc.sections]
        cols = ['Sl. No', 'Reg Number', 'Name of the student', 'Midterm Percentage', 'Progress']
        
        if not students:
            if hasattr(self, 'subject'): doc.add_paragraph(f"Course: {str(self.subject).upper()}", style='Heading 3')
            t = doc.add_table(rows=1, cols=len(cols)); t.style = 'Table Grid'
            for j, c in enumerate(cols): self.set_cell_properties(t.cell(0, j), c, bold=True)
            for i in range(8):
                rc = t.add_row().cells; self.set_cell_properties(rc[0], str(i+1), font_name=self.BODY_FONT)
                if i == 3: self.set_cell_properties(rc[2], "NIL", font_name=self.BODY_FONT, align='CENTER')
            self.add_signature_line(doc); return doc
        
        group = pd.DataFrame(students)
        doc.add_paragraph(f"Course: {str(group['Subject Name'].iloc[0]).upper()}", style='Heading 3')
        t = doc.add_table(rows=1, cols=len(cols)); t.style = 'Table Grid'
        for j, c in enumerate(cols): self.set_cell_properties(t.cell(0, j), c, bold=True)
        for idx, row in group.reset_index(drop=True).iterrows():
            rc = t.add_row().cells; self.set_cell_properties(rc[0], str(idx+1))
            self.set_cell_properties(rc[1], row['Register Number of the Student'])
            self.set_cell_properties(rc[2], row['Student Name'])
            self.set_cell_properties(rc[3], f"{row['MidtermPercentage']:.2f}")
            self.set_cell_properties(rc[4], row['Outcome (Based on clearance in end-semester or makeup exam)'])
        self.add_signature_line(doc); return doc

class Format1And2DocxFormatter(BaseFormatter):
    def format(self, students, st, ft):
        doc = Document(); [setattr(sec, 'top_margin', Inches(0.5)) for sec in doc.sections]
        for i, student in enumerate(students): 
            self._create_format1_content(doc, student, st, ft); doc.add_page_break()
            self._create_format2_content(doc, student); (i < len(students)-1) and doc.add_page_break()
        return doc

# --- WRITERS ---
class DocxWriter:
    def write(self, doc, out, **kwargs): doc.save(out)
class PdfWriter:
    def write(self, doc, out, sign_info=None, format_choice=None):
        with tempfile.TemporaryDirectory() as td:
            temp_docx = os.path.join(td, "temp.docx")
            doc.save(temp_docx)
            subprocess.run([get_libreoffice_command(), '--headless', '--convert-to', 'pdf', '--outdir', td, temp_docx], check=True, stdout=subprocess.DEVNULL)
            if os.path.exists(os.path.join(td, "temp.pdf")):
                import shutil; shutil.move(os.path.join(td, "temp.pdf"), out)
                if sign_info and sign_info.get('should_sign') and format_choice in ['1','2','4','5']:
                    sign_pdf(out, sign_info['key_path'], sign_info['cert_path'], sign_info['image_path'], sign_info['password'])

# --- CONTROLLER ---
class ReportController:
    def __init__(self, excel_path, cgpa_path, format_choice, learner_type, slow_thresh, advanced_thresh, output_type, semester, sign_info, common_comment, grade_path=None, faculty_name=None):
        self.excel_path = excel_path; self.cgpa_path = cgpa_path; self.grade_path = grade_path; self.format_choice = format_choice; self.learner_type = learner_type
        self.slow_threshold = slow_thresh; self.advanced_threshold = advanced_thresh; self.output_type = output_type; self.semester = semester.lower().strip()
        self.sign_info = sign_info; self.common_comment = common_comment; self.faculty_name = faculty_name
        self.reader = DataReader(); self.processor = StudentDataProcessor()
        self.writer = DocxWriter() if output_type == 'word' else PdfWriter()

    def run(self):
        all_data, self.subject = self.reader.read_data(self.excel_path)
        if not all_data: return None
        cg_map = self.reader.read_cgpa_map(self.cgpa_path)
        grade_map = self.reader.read_grade_map(self.grade_path, course_code=self.subject)
        processed = self.processor.process_data(all_data, self.subject, self.semester, self.common_comment, cg_map, grade_map)
        filtered = self.processor.filter_students(processed, self.learner_type, self.slow_threshold, self.advanced_threshold)
        
        act_f = '3' if not filtered else self.format_choice
        ds = datetime.now().strftime('%d_%m_%y'); sn = self.semester.upper()
        sub_dir = re.sub(r'[\\/*?:"<>|]', "", self.subject.replace(' ', '_'))
        od = os.path.join("Learner_Monitor_Reports", f"{self.learner_type.title()}_Learners", f"Semester_{sn}", sub_dir)
        os.makedirs(od, exist_ok=True)

        if act_f == '5' and filtered: return self._generate_all_formats(filtered, od, ds, sn, sub_dir)
        
        fmt = {'1': Format1DocxFormatter(), '2': Format2DocxFormatter(), '3': Format3DocxFormatter(), '4': Format1And2DocxFormatter()}.get(act_f)
        for attr in ['signature_image_path', 'faculty_name', 'learner_type', 'subject', 'semester']:
            val = getattr(self, attr) if hasattr(self, attr) else self.sign_info.get('image_path') if attr == 'signature_image_path' else None
            setattr(fmt, attr, val)
        
        ext = 'docx' if self.output_type == 'word' else 'pdf'
        lbl = {'1':'Format1', '2':'Format2', '3':'Summary', '4':'Combined'}.get(act_f, "Report") if filtered else "Empty_Summary"
        out_p = os.path.join(od, f'{sub_dir}_{sn}_{self.learner_type.title()}Learner_{lbl}_{ds}.{ext}')
        self.writer.write(fmt.format(filtered, self.slow_threshold, self.advanced_threshold), out_p, sign_info=self.sign_info, format_choice=act_f)
        return out_p

    def _generate_all_formats(self, students, od, ds, sn, sub_dir):
        ext = 'docx' if self.output_type == 'word' else 'pdf'
        res = []
        for f_class, lbl, fc in [(Format1And2DocxFormatter, 'Combined', '4'), (Format3DocxFormatter, 'Summary', '3')]:
            fmt = f_class()
            fmt.signature_image_path, fmt.faculty_name, fmt.learner_type, fmt.subject, fmt.semester = self.sign_info.get('image_path'), self.faculty_name, self.learner_type, self.subject, self.semester
            p = os.path.join(od, f'{sub_dir}_{sn}_{self.learner_type.title()}Learner_{lbl}_{ds}.{ext}')
            self.writer.write(fmt.format(students, self.slow_threshold, self.advanced_threshold), p, sign_info=self.sign_info, format_choice=fc)
            res.append(p)
        return res